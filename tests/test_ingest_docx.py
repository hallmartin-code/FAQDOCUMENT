"""DOCX ingestion.

A Word document has no slides and no fixed pages, so the contract under test is that it
is split at its headings into numbered sections, and that a citation of section N points
somewhere a reader can actually find.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from deckpager.errors import IngestError, UnsupportedFormatError
from deckpager.extract.prompts import build_user_blocks, slide_text, unit_label
from deckpager.ingest.docx import MAX_SECTION_CHARS, load_docx
from deckpager.ingest.router import load_deck


def _write(path: Path, blocks: list[tuple[str, str]]) -> Path:
    """Build a .docx from (style, text) pairs. 'table' rows are pipe-separated."""
    document = Document()
    for style, text in blocks:
        if style == "table":
            rows = [row.split("|") for row in text.split("\n")]
            table = document.add_table(rows=len(rows), cols=len(rows[0]))
            for r, row in enumerate(rows):
                for c, cell in enumerate(row):
                    table.cell(r, c).text = cell.strip()
        else:
            document.add_paragraph(text, style=style or None)
    document.save(str(path))
    return path


@pytest.fixture
def memo(tmp_path: Path) -> Path:
    return _write(
        tmp_path / "memo.docx",
        [
            ("Title", "Northwind Robotics"),
            ("", "Warehouse automation for mid-market distributors."),
            ("Heading 1", "The problem"),
            ("", "Pickers walk 12 miles a shift and turnover runs at 90% a year."),
            ("Heading 1", "Traction"),
            ("", "Three paid pilots signed in Q1."),
            ("table", "Metric|2025|2026\nARR|$400K|$1.2M"),
        ],
    )


class TestSections:
    def test_headings_become_numbered_sections(self, memo: Path) -> None:
        deck = load_docx(memo)
        assert [s.title for s in deck.slides] == [
            "Northwind Robotics",
            "The problem",
            "Traction",
        ]
        assert [s.index for s in deck.slides] == [1, 2, 3]

    def test_the_format_is_recorded(self, memo: Path) -> None:
        assert load_docx(memo).source_format == "docx"

    def test_prose_lands_under_its_own_heading(self, memo: Path) -> None:
        deck = load_docx(memo)
        assert "12 miles a shift" in deck.slides[1].text
        assert "12 miles a shift" not in deck.slides[0].text

    def test_tables_are_flattened_in_place(self, memo: Path) -> None:
        """A table must stay with the prose it belongs to, not be appended at the end."""
        deck = load_docx(memo)
        assert "ARR: $400K | $1.2M" in deck.slides[2].text

    def test_no_page_images_are_produced(self, memo: Path) -> None:
        assert all(s.asset is None for s in load_docx(memo).slides)


class TestDegenerateDocuments:
    def test_a_document_without_headings_is_still_split(self, tmp_path: Path) -> None:
        long_prose = [("", "Sentence number %d in a long unstructured memo." % i) for i in range(220)]
        deck = load_docx(_write(tmp_path / "flat.docx", long_prose))
        assert deck.slide_count > 1
        assert all(len(s.text) <= MAX_SECTION_CHARS * 1.5 for s in deck.slides)
        assert any("no headings" in w for w in deck.warnings)

    def test_an_empty_document_fails_with_an_actionable_message(self, tmp_path: Path) -> None:
        with pytest.raises(IngestError) as exc:
            load_docx(_write(tmp_path / "empty.docx", [("", "   ")]))
        assert "no readable text" in str(exc.value)

    def test_inline_images_are_declared_not_silently_dropped(self, tmp_path: Path) -> None:
        path = tmp_path / "withimage.docx"
        document = Document()
        document.add_heading("Product", level=1)
        document.add_paragraph("The device is shown below.")
        png = (
            b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x06"
            b"\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01\x00\x00\x05"
            b"\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82"
        )
        (tmp_path / "dot.png").write_bytes(png)
        document.add_picture(str(tmp_path / "dot.png"))
        document.save(str(path))

        deck = load_docx(path)
        assert any("inline image" in w for w in deck.warnings)


class TestRouting:
    def test_the_router_dispatches_docx(self, memo: Path) -> None:
        assert load_deck(memo, want_images=False).source_format == "docx"

    def test_a_docx_renamed_to_pptx_is_refused(self, memo: Path, tmp_path: Path) -> None:
        """Both are ZIPs; only the part names tell them apart."""
        renamed = tmp_path / "disguised.pptx"
        renamed.write_bytes(memo.read_bytes())
        with pytest.raises(UnsupportedFormatError) as exc:
            load_deck(renamed, want_images=False)
        assert "DOCX" in str(exc.value).upper()


class TestPromptFraming:
    def test_sections_are_not_called_slides(self, memo: Path) -> None:
        deck = load_docx(memo)
        assert unit_label(deck) == "SECTION"
        payload = slide_text(deck)
        assert "--- SECTION 1 ---" in payload
        assert "SLIDE" not in payload

    def test_the_instruction_says_what_the_numbers_mean(self, memo: Path) -> None:
        text = build_user_blocks(load_docx(memo))[-1]["text"]
        assert "not a slide deck" in text
        assert "source_slides" in text


class TestBoldHeadings:
    """Most real memos carry no outline styles — they bold a line and move on."""

    def _bold_doc(self, path: Path, headings: list[str]) -> Path:
        document = Document()
        for heading in headings:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(heading)
            run.bold = True
            document.add_paragraph(f"Body text belonging to {heading}, stated plainly.")
        document.save(str(path))
        return path

    def test_bold_lines_are_treated_as_headings(self, tmp_path: Path) -> None:
        path = self._bold_doc(tmp_path / "memo.docx", ["Target Markets", "Use of Funds"])
        deck = load_docx(path)
        assert [s.title for s in deck.slides] == ["Target Markets", "Use of Funds"]
        assert any("bold single-line paragraphs" in w for w in deck.warnings)

    def test_a_bolded_sentence_is_not_a_heading(self, tmp_path: Path) -> None:
        """Emphasis inside prose must not open a section."""
        path = tmp_path / "emphasis.docx"
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("We grew revenue by 300% last year and expect to double again.").bold = True
        document.add_paragraph("Supporting detail follows here.")
        document.save(str(path))
        assert load_docx(path).slide_count == 1

    def test_styled_headings_win_over_the_bold_heuristic(self, tmp_path: Path) -> None:
        """The heuristic is a fallback; a properly styled document never reaches it."""
        path = tmp_path / "styled.docx"
        document = Document()
        document.add_heading("Real Heading", level=1)
        document.add_paragraph("Prose under the real heading.")
        emphasised = document.add_paragraph()
        emphasised.add_run("Bold Line").bold = True
        document.add_heading("Second Heading", level=1)
        document.add_paragraph("More prose.")
        document.save(str(path))

        deck = load_docx(path)
        assert [s.title for s in deck.slides] == ["Real Heading", "Second Heading"]
        assert "Bold Line" in deck.slides[0].text
        assert not any("bold single-line" in w for w in deck.warnings)
