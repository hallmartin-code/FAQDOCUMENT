"""DOCX ingestion.

A Word document has no slides, and Word has no fixed pages until something renders it.
So the unit of citation here is the **section**: the document is split at its headings,
and each section becomes a `Slide` with a 1-based index. Every `source_slides` citation
the model produces then points at a section of the document, and the reader can find it
by its heading rather than by counting pages that do not exist.

Text only. There is no page rasterization: without LibreOffice there is nothing to render
a .docx into, and unlike a deck, a memo or an executive summary carries its argument in
prose rather than in pictures. Inline images are noted, not sent.
"""

from __future__ import annotations

from collections.abc import Iterator
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from deckpager.errors import IngestError
from deckpager.ingest.models import Deck, Slide, normalize_text
from deckpager.ingest.pdf import first_line_title, flatten_table

#: Styles that open a new section. "Title" is included because a document whose only
#: structural marker is its title should still produce a titled first section.
_HEADING_PREFIXES = ("Heading", "Title", "Subtitle")

#: A section longer than this is split across continuation sections. An unstructured
#: document — no headings at all — would otherwise arrive as one giant section, and a
#: citation of "section 1" for a 12-page memo tells a reader nothing.
MAX_SECTION_CHARS = 3500

#: Below this, a trailing fragment is folded back into the previous section rather than
#: standing as its own; a section holding one orphaned sentence is noise to cite.
MIN_SECTION_CHARS = 40

#: A bold line longer than this is a sentence someone emphasised, not a heading.
PSEUDO_HEADING_MAX_CHARS = 90

#: Bounds on the pseudo-heading pass. Under two and it achieved nothing; over this and the
#: document is one where bold means emphasis, and every emphasised phrase would open a
#: section. Either way the character split is the more honest fallback.
PSEUDO_HEADING_MAX_SECTIONS = 40


def _is_heading(paragraph: Paragraph) -> bool:
    """Whether a paragraph opens a new section, by its style."""
    style = paragraph.style
    name = getattr(style, "name", "") or ""
    return name.startswith(_HEADING_PREFIXES)


def _looks_like_heading(paragraph: Paragraph) -> bool:
    """Whether a paragraph reads as a heading despite carrying no heading style.

    Real memos and executive summaries are rarely written with Word's outline styles.
    People bold a short line and move on, which leaves a document that is plainly
    structured to a reader and completely flat to `paragraph.style`. Without this, a
    12-page summary gets chopped at arbitrary character counts and every citation points
    at a section boundary the reader cannot see.
    """
    text = paragraph.text.strip()
    if not text or len(text) > PSEUDO_HEADING_MAX_CHARS:
        return False
    # A heading may end in a colon; ending in sentence punctuation means it is prose.
    if text[-1] in ".,;!?":
        return False
    runs = [run for run in paragraph.runs if run.text.strip()]
    if not runs:
        return False
    return all(bool(run.bold) for run in runs)


def _iter_blocks(document: Any) -> Iterator[Paragraph | Table]:
    """Yield paragraphs and tables in document order.

    python-docx exposes `.paragraphs` and `.tables` as separate sequences, which loses
    the interleaving — a table's position relative to the prose around it is exactly the
    context that makes it readable, so the body XML is walked directly instead.
    """
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _table_lines(table: Table) -> list[str]:
    """Flatten a table to one line per row, as the PDF and PPTX paths do."""
    rows: list[list[str | None]] = [[cell.text for cell in row.cells] for row in table.rows]
    return flatten_table(rows)


def _split_long(title: str | None, lines: list[str]) -> list[tuple[str | None, str]]:
    """Break an over-long section on paragraph boundaries, preserving the heading."""
    chunks: list[tuple[str | None, str]] = []
    current: list[str] = []
    size = 0
    for line in lines:
        # A single paragraph longer than the budget is kept whole: splitting mid-sentence
        # would hand the model a fragment it cannot quote accurately.
        if current and size + len(line) > MAX_SECTION_CHARS:
            chunks.append((title, "\n".join(current)))
            current, size = [], 0
        current.append(line)
        size += len(line) + 1
    if current:
        chunks.append((title, "\n".join(current)))
    if len(chunks) > 1:
        chunks = [
            (t if i == 0 else f"{t} (cont.)" if t else None, body)
            for i, (t, body) in enumerate(chunks)
        ]
    return chunks


def _split_into_sections(
    blocks: list[Paragraph | Table], *, styled_only: bool
) -> list[tuple[str | None, list[str]]]:
    """Walk the document once, opening a new section at each heading."""
    sections: list[tuple[str | None, list[str]]] = []
    title: str | None = None
    lines: list[str] = []

    def close() -> None:
        if any(line.strip() for line in lines):
            sections.append((title, list(lines)))

    for block in blocks:
        if isinstance(block, Table):
            lines.extend(_table_lines(block))
            continue
        text = block.text.strip()
        heading = _is_heading(block) or (not styled_only and _looks_like_heading(block))
        if heading and text:
            close()
            title, lines = text[:200], []
            continue
        if text:
            lines.append(text)
    close()
    return sections


def load_docx(path: Path, *, want_images: bool = True) -> Deck:
    """Read a DOCX into a Deck, one section per heading.

    `want_images` is accepted for signature parity with the other loaders and ignored:
    there is no page image to produce.
    """
    try:
        document = Document(str(path))
    except Exception as exc:  # python-docx raises package-specific errors for bad files
        raise IngestError(f"{path.name} is not a readable DOCX: {exc}") from exc

    warnings: list[str] = []
    blocks = list(_iter_blocks(document))
    inline_images = sum(
        1
        for b in blocks
        if isinstance(b, Paragraph) and b._element.findall(f".//{qn('a:blip')}")
    )

    sections = _split_into_sections(blocks, styled_only=True)

    # Second pass: a document with no outline styles is not necessarily unstructured.
    pseudo_used = False
    if len(sections) <= 1:
        candidate = _split_into_sections(blocks, styled_only=False)
        if 1 < len(candidate) <= PSEUDO_HEADING_MAX_SECTIONS:
            sections = candidate
            pseudo_used = True

    if not sections:
        raise IngestError(
            f"{path.name} contains no readable text. If it is a scanned document, "
            f"export a text-bearing PDF instead."
        )

    # Fold a runt trailing section into its predecessor — but never a titled one. A
    # heading is a citation target the reader can find, however short its body; folding
    # it away deletes a place the model is allowed to point at.
    if (
        len(sections) > 1
        and sections[-1][0] is None
        and sum(len(x) for x in sections[-1][1]) < MIN_SECTION_CHARS
    ):
        tail_title, tail_lines = sections.pop()
        head_title, head_lines = sections[-1]
        sections[-1] = (head_title, [*head_lines, *([tail_title] if tail_title else []), *tail_lines])

    slides: list[Slide] = []
    for section_title, section_lines in sections:
        for chunk_title, body in _split_long(section_title, section_lines):
            normalized = normalize_text(body)
            if not normalized:
                continue
            slides.append(
                Slide(
                    index=len(slides) + 1,
                    title=chunk_title or first_line_title(normalized),
                    text=normalized,
                    speaker_notes=None,
                    asset=None,
                    has_chart=False,
                )
            )

    if not slides:
        raise IngestError(f"{path.name} contains no readable text.")

    if len(sections) == 1 and len(slides) > 1:
        warnings.append(
            f"{path.name} has no headings; it was split into {len(slides)} sections of "
            f"roughly {MAX_SECTION_CHARS} characters so citations can point somewhere. "
            f"Section numbers will not match anything visible in the document."
        )
    elif pseudo_used:
        warnings.append(
            f"{path.name} uses no heading styles; its {len(sections)} sections were found "
            f"from bold single-line paragraphs. Section numbers follow those lines."
        )
    if inline_images:
        warnings.append(
            f"{path.name} contains {inline_images} inline image(s), which are not sent: "
            f"a DOCX has no page rendering to attach. Anything stated only inside a chart "
            f"or screenshot will be missed."
        )

    return Deck(
        source_path=path,
        source_format="docx",
        slides=slides,
        raw_pdf_b64=None,
        warnings=warnings,
    )
